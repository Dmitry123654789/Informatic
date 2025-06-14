import base64

from flask import Flask, render_template
import os
from docx import Document
from data.db_session import global_init
app = Flask(__name__)
my_dir = os.path.dirname(__file__)
NAMES_TESTS = sorted(os.listdir(os.path.join(my_dir, 'tests_9')))


@app.route('/')
def index():
    return render_template('index.html', names=NAMES_TESTS)


@app.route('/tests/<test_id>')
def show_test(test_id):
    doc = Document(os.path.join(my_dir, f'tests_9/{NAMES_TESTS[int(test_id) - 1]}'))
    new_lst = []
    for x in [x.text for x in doc.paragraphs]:
        new_lst += [y.strip() for y in x.split('\n')]

    images = []
    for rel in doc.part._rels:
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            img_base64 = base64.b64encode(img_data).decode('utf-8')
            s = rel.target_ref.split('.')
            images.append((img_base64, int(s[0][11:])))
    images.sort(key=lambda x: x[1])
    i = 0
    for n, x in enumerate(new_lst):
        if x == 'img':
            try:
                new_lst[n] = 'img' + images[i][0]
                i += 1
            except IndexError:
                break

    return render_template('show_test.html', strings=new_lst, images=images)


@app.route('/all_tests')
def show_all_test():
    final_lst = []
    for name in NAMES_TESTS:
        doc = Document(os.path.join(my_dir, f'tests_9/{name}'))
        new_lst = []
        for x in [x.text for x in doc.paragraphs]:
            new_lst += [y.strip() for y in x.split('\n')]

        images = []
        for rel in doc.part._rels:
            rel = doc.part._rels[rel]
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                s = rel.target_ref.split('.')
                print(s)

        images.sort(key=lambda x: x[1])
        i = 0
        for n, x in enumerate(new_lst):
            if x == 'img':
                try:
                    new_lst[n] = 'img' + images[i][0]
                    i += 1
                except IndexError:
                    break

        final_lst.append('')
        final_lst.append('')
        final_lst += new_lst

    return render_template('show_test.html', strings=final_lst)


if __name__ == '__main__':
    global_init('db/informatic_questions.db')
    app.run()
